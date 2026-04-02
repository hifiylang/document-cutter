from __future__ import annotations

"""Office 文档解析器。"""

import io
import os
import re
import shutil
import subprocess
import tempfile
import uuid

from docx import Document as DocxDocument
from openpyxl import load_workbook
import xlrd

from app.models.schemas import DocumentNode
from app.services.parsers.base import BaseParser


class DocxParser(BaseParser):
    """解析 Word `.docx` 中的标题、正文与表格。"""

    def parse(self, file_bytes: bytes, filename: str) -> list[DocumentNode]:
        document = DocxDocument(io.BytesIO(file_bytes))
        nodes: list[DocumentNode] = []
        for paragraph in document.paragraphs:
            text = (paragraph.text or "").strip()
            if not text:
                continue
            style_name = (paragraph.style.name or "").lower() if paragraph.style else ""
            if style_name.startswith("heading"):
                digits = "".join(ch for ch in style_name if ch.isdigit())
                level = int(digits) if digits else 1
                nodes.append(
                    DocumentNode(
                        node_id=str(uuid.uuid4()),
                        node_type="title",
                        level=level,
                        text=text,
                        source_meta={"parser_strategy": "docx_heading"},
                    )
                )
            else:
                node_type = "list" if re.match(r"^\s*(?:[-*+]|\d+[.)])\s+", text) else "paragraph"
                nodes.append(
                    DocumentNode(
                        node_id=str(uuid.uuid4()),
                        node_type=node_type,
                        text=text,
                        source_meta={"parser_strategy": "docx_paragraph"},
                    )
                )

        for table in document.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    rows.append(" | ".join(cells))
            table_text = "\n".join(rows).strip()
            if table_text:
                nodes.append(
                    DocumentNode(
                        node_id=str(uuid.uuid4()),
                        node_type="table",
                        text=table_text,
                        source_meta={"parser_strategy": "docx_table"},
                    )
                )
        return nodes


class DocParser(BaseParser):
    """解析旧版 Word `.doc` 文档。"""

    heading_re = re.compile(
        r"^(?:第[一二三四五六七八九十百千万\d]+[章节部分篇]|[一二三四五六七八九十]+[、.]|\d+(?:\.\d+){0,3}[、.]?)?\s*\S+$"
    )
    list_re = re.compile(r"^\s*(?:[-*+•●]|\d+[.)]|（\d+）)\s+")

    def parse(self, file_bytes: bytes, filename: str) -> list[DocumentNode]:
        text = self._extract_text(file_bytes)
        lines = [line.strip() for line in text.splitlines()]
        nodes: list[DocumentNode] = []
        paragraph_buf: list[str] = []
        list_buf: list[str] = []

        def flush_paragraph() -> None:
            if not paragraph_buf:
                return
            body = " ".join(paragraph_buf).strip()
            if body:
                nodes.append(
                    DocumentNode(
                        node_id=str(uuid.uuid4()),
                        node_type="paragraph",
                        text=body,
                        source_meta={"parser_strategy": "doc_antiword"},
                    )
                )
            paragraph_buf.clear()

        def flush_list() -> None:
            if not list_buf:
                return
            body = "\n".join(list_buf).strip()
            if body:
                nodes.append(
                    DocumentNode(
                        node_id=str(uuid.uuid4()),
                        node_type="list",
                        text=body,
                        source_meta={"parser_strategy": "doc_antiword"},
                    )
                )
            list_buf.clear()

        for line in lines:
            stripped = line.strip()
            if not stripped:
                flush_paragraph()
                flush_list()
                continue

            if self.list_re.match(stripped):
                flush_paragraph()
                list_buf.append(stripped)
                continue

            if self._looks_like_heading(stripped):
                flush_paragraph()
                flush_list()
                nodes.append(
                    DocumentNode(
                        node_id=str(uuid.uuid4()),
                        node_type="title",
                        level=1,
                        text=stripped,
                        source_meta={"parser_strategy": "doc_antiword"},
                    )
                )
                continue

            flush_list()
            paragraph_buf.append(stripped)

        flush_paragraph()
        flush_list()
        return nodes

    def _extract_text(self, file_bytes: bytes) -> str:
        antiword = shutil.which("antiword")
        if not antiword:
            raise RuntimeError("antiword is required to parse .doc files")

        with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as temp_file:
            temp_file.write(file_bytes)
            temp_path = temp_file.name

        try:
            completed = subprocess.run(
                [antiword, temp_path],
                check=True,
                capture_output=True,
                text=True,
                encoding="utf-8",
                errors="ignore",
            )
            return completed.stdout or ""
        except subprocess.CalledProcessError as exc:
            stderr = (exc.stderr or "").strip()
            raise RuntimeError(f"failed to parse .doc file with antiword: {stderr or 'unknown error'}") from exc
        finally:
            try:
                os.unlink(temp_path)
            except OSError:
                pass

    def _looks_like_heading(self, text: str) -> bool:
        if len(text) > 24:
            return False
        if re.search(r"[。！？!?；;，,]", text):
            return False
        if self.list_re.match(text):
            return False
        return bool(self.heading_re.match(text))


class XlsxParser(BaseParser):
    """按 sheet 提取 XLSX 内容。"""

    def parse(self, file_bytes: bytes, filename: str) -> list[DocumentNode]:
        workbook = load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True)
        nodes: list[DocumentNode] = []
        for sheet in workbook.worksheets:
            nodes.append(
                DocumentNode(
                    node_id=str(uuid.uuid4()),
                    node_type="title",
                    level=1,
                    text=sheet.title,
                    source_meta={"sheet_name": sheet.title, "parser_strategy": "xlsx_sheet"},
                )
            )
            rows = []
            for row in sheet.iter_rows(values_only=True):
                values = ["" if value is None else str(value).strip() for value in row]
                if any(values):
                    rows.append(" | ".join(values))
            table_text = "\n".join(rows).strip()
            if table_text:
                nodes.append(
                    DocumentNode(
                        node_id=str(uuid.uuid4()),
                        node_type="table",
                        text=table_text,
                        source_meta={"sheet_name": sheet.title, "parser_strategy": "xlsx_table"},
                    )
                )
        return nodes


class XlsParser(BaseParser):
    """按 sheet 提取 XLS 内容。"""

    def parse(self, file_bytes: bytes, filename: str) -> list[DocumentNode]:
        workbook = xlrd.open_workbook(file_contents=file_bytes)
        nodes: list[DocumentNode] = []
        for sheet in workbook.sheets():
            nodes.append(
                DocumentNode(
                    node_id=str(uuid.uuid4()),
                    node_type="title",
                    level=1,
                    text=sheet.name,
                    source_meta={"sheet_name": sheet.name, "parser_strategy": "xls_sheet"},
                )
            )
            rows = []
            for row_index in range(sheet.nrows):
                values = []
                for value in sheet.row_values(row_index):
                    text = "" if value is None else str(value).strip()
                    if text.endswith(".0"):
                        text = text[:-2]
                    values.append(text)
                if any(values):
                    rows.append(" | ".join(values))
            table_text = "\n".join(rows).strip()
            if table_text:
                nodes.append(
                    DocumentNode(
                        node_id=str(uuid.uuid4()),
                        node_type="table",
                        text=table_text,
                        source_meta={"sheet_name": sheet.name, "parser_strategy": "xls_table"},
                    )
                )
        return nodes

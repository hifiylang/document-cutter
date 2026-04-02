from __future__ import annotations
"""把不同格式的文档解析成统一的 DocumentNode 结构。"""

import io
import logging
import re
import uuid
from abc import ABC, abstractmethod
from collections import Counter
from dataclasses import dataclass
from pathlib import Path

from docx import Document as DocxDocument
from openpyxl import load_workbook
from pypdf import PdfReader
import xlrd

from app.core.metrics import (
    PDF_IMAGE_REGION_DETECTED,
    PDF_IMAGE_REGION_VISION_ERROR,
    PDF_IMAGE_REGION_VISION_SUCCESS,
)
from app.models.schemas import DocumentNode
from app.services.vision import VisualDocumentAnalyzer

try:
    import fitz
except Exception:  # pragma: no cover
    fitz = None  # type: ignore


logger = logging.getLogger(__name__)

IMAGE_SUFFIXES = {".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tif", ".tiff"}


@dataclass
class PdfImageRegion:
    page_no: int
    bbox: list[float]
    image_region_id: str
    order: int
    parser_strategy: str = "pdf_image_region"
    layout_role: str = "image_region"


class BaseParser(ABC):
    @abstractmethod
    def parse(self, file_bytes: bytes, filename: str) -> list[DocumentNode]:
        """解析原始文件内容并返回标准节点。"""
        raise NotImplementedError


class TxtMarkdownParser(BaseParser):
    heading_re = re.compile(r"^(#{1,6})\s+(.*)$")
    list_re = re.compile(r"^(\s*(?:[-*+]|\d+[.)]))\s+(.+)$")

    def parse(self, file_bytes: bytes, filename: str) -> list[DocumentNode]:
        """解析 TXT 和 Markdown，保留标题、列表、表格和段落结构。"""
        text = file_bytes.decode("utf-8", errors="ignore")
        lines = text.splitlines(keepends=True)
        nodes: list[DocumentNode] = []
        paragraph_buf: list[str] = []
        paragraph_start: int | None = None
        list_buf: list[str] = []
        list_start: int | None = None
        table_buf: list[str] = []
        table_start: int | None = None
        cursor = 0

        def flush_paragraph(end: int) -> None:
            nonlocal paragraph_start
            if not paragraph_buf or paragraph_start is None:
                return
            body = "".join(paragraph_buf).strip()
            if body:
                body_start = text.find(body, paragraph_start, end)
                nodes.append(
                    DocumentNode(
                        node_id=str(uuid.uuid4()),
                        node_type="paragraph",
                        text=body,
                        source_meta={
                            "char_start": body_start,
                            "char_end": body_start + len(body),
                            "parser_strategy": "markdown_text",
                        },
                    )
                )
            paragraph_buf.clear()
            paragraph_start = None

        def flush_list(end: int) -> None:
            nonlocal list_start
            if not list_buf or list_start is None:
                return
            body = "".join(list_buf).strip()
            if body:
                body_start = text.find(body, list_start, end)
                nodes.append(
                    DocumentNode(
                        node_id=str(uuid.uuid4()),
                        node_type="list",
                        text=body,
                        source_meta={
                            "char_start": body_start,
                            "char_end": body_start + len(body),
                            "parser_strategy": "markdown_text",
                        },
                    )
                )
            list_buf.clear()
            list_start = None

        def flush_table(end: int) -> None:
            nonlocal table_start
            if not table_buf or table_start is None:
                return
            body = "".join(table_buf).strip()
            if body:
                body_start = text.find(body, table_start, end)
                nodes.append(
                    DocumentNode(
                        node_id=str(uuid.uuid4()),
                        node_type="table",
                        text=body,
                        source_meta={
                            "char_start": body_start,
                            "char_end": body_start + len(body),
                            "parser_strategy": "markdown_table",
                        },
                    )
                )
            table_buf.clear()
            table_start = None

        for raw in lines:
            line_start = cursor
            cursor += len(raw)
            line = raw.rstrip("\r\n")
            stripped = line.strip()
            heading_match = self.heading_re.match(stripped)
            is_table_row = "|" in stripped and stripped.count("|") >= 2
            is_list_row = bool(self.list_re.match(stripped))

            if heading_match:
                flush_paragraph(line_start)
                flush_list(line_start)
                flush_table(line_start)
                marks, title = heading_match.groups()
                body_start = text.find(title.strip(), line_start, cursor)
                nodes.append(
                    DocumentNode(
                        node_id=str(uuid.uuid4()),
                        node_type="title",
                        level=len(marks),
                        text=title.strip(),
                        source_meta={
                            "char_start": body_start,
                            "char_end": body_start + len(title.strip()),
                            "parser_strategy": "markdown_title",
                        },
                    )
                )
                continue

            if is_table_row:
                flush_paragraph(line_start)
                flush_list(line_start)
                if table_start is None:
                    table_start = line_start
                table_buf.append(raw)
                continue

            if table_buf:
                flush_table(line_start)

            if is_list_row:
                flush_paragraph(line_start)
                if list_start is None:
                    list_start = line_start
                list_buf.append(raw)
                continue

            if list_buf:
                flush_list(line_start)

            if not stripped:
                flush_paragraph(line_start)
                continue

            if paragraph_start is None:
                paragraph_start = line_start
            paragraph_buf.append(raw)

        flush_paragraph(len(text))
        flush_list(len(text))
        flush_table(len(text))
        return nodes


class DocxParser(BaseParser):
    def parse(self, file_bytes: bytes, filename: str) -> list[DocumentNode]:
        """解析 Word 文档中的标题、正文和表格。"""
        doc = DocxDocument(io.BytesIO(file_bytes))
        nodes: list[DocumentNode] = []
        for p in doc.paragraphs:
            text = (p.text or "").strip()
            if not text:
                continue
            style_name = (p.style.name or "").lower() if p.style else ""
            if style_name.startswith("heading"):
                digits = "".join(ch for ch in style_name if ch.isdigit())
                level = int(digits) if digits else 1
                nodes.append(
                    DocumentNode(
                        node_id=str(uuid.uuid4()),
                        node_type="title",
                        level=level,
                        text=text,
                        source_meta={"parser_strategy": "docx_heading"},
                    )
                )
            else:
                node_type = "list" if re.match(r"^\s*(?:[-*+]|\d+[.)])\s+", text) else "paragraph"
                nodes.append(
                    DocumentNode(
                        node_id=str(uuid.uuid4()),
                        node_type=node_type,
                        text=text,
                        source_meta={"parser_strategy": "docx_paragraph"},
                    )
                )

        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    rows.append(" | ".join(cells))
            table_text = "\n".join(rows).strip()
            if table_text:
                nodes.append(
                    DocumentNode(
                        node_id=str(uuid.uuid4()),
                        node_type="table",
                        text=table_text,
                        source_meta={"parser_strategy": "docx_table"},
                    )
                )
        return nodes


class PdfParser(BaseParser):
    page_number_re = re.compile(r"^(?:page\s*\d+|第\s*\d+\s*页|\d+\s*/\s*\d+)$", re.IGNORECASE)

    def __init__(self) -> None:
        self.visual_analyzer = VisualDocumentAnalyzer()

    def parse(self, file_bytes: bytes, filename: str) -> list[DocumentNode]:
        """优先走 PDF 版面解析，再按需补局部图片视觉解析。"""
        nodes: list[DocumentNode] = []
        if fitz is not None:
            nodes = self._extract_with_pymupdf(file_bytes, filename)
        if not nodes:
            nodes = self._extract_with_pypdf(file_bytes)
        return self._remove_repeated_page_noise(nodes)

    def _extract_with_pymupdf(self, file_bytes: bytes, filename: str) -> list[DocumentNode]:
        document = fitz.open(stream=file_bytes, filetype="pdf")
        nodes: list[DocumentNode] = []
        for page_index, page in enumerate(document, start=1):
            page_height = float(page.rect.height or 1)
            text_blocks = page.get_text("blocks")
            page_tables = self._extract_tables_from_page(page, page_index)
            table_bboxes = [table.source_meta.get("bbox") for table in page_tables if table.source_meta.get("bbox")]
            page_regions = self._extract_image_regions(page, page_index, table_bboxes)
            page_text_nodes = self._extract_text_nodes(page_index, page_height, text_blocks, table_bboxes)
            page_image_nodes = self._extract_image_nodes(page, page_index, filename, page_regions)
            page_nodes = self._assemble_page_nodes(page_text_nodes, page_tables, page_image_nodes)
            logger.info(
                "pdf page=%s image_regions=%s image_nodes=%s text_nodes=%s table_nodes=%s",
                page_index,
                len(page_regions),
                len(page_image_nodes),
                len(page_text_nodes),
                len(page_tables),
            )
            nodes.extend(page_nodes)
        return nodes

    def _extract_text_nodes(
        self,
        page_index: int,
        page_height: float,
        blocks,
        table_bboxes: list[object],
    ) -> list[DocumentNode]:
        nodes: list[DocumentNode] = []
        for block in self._sort_blocks_by_columns(blocks):
            x0, y0, x1, y1, text, *_ = block
            bbox = [float(x0), float(y0), float(x1), float(y1)]
            if self._overlaps_any(bbox, table_bboxes):
                continue
            cleaned = self._clean_pdf_text(text)
            if not cleaned:
                continue
            node_type, level = self._infer_pdf_node_type(cleaned)
            nodes.append(
                DocumentNode(
                    node_id=str(uuid.uuid4()),
                    node_type=node_type,
                    level=level,
                    text=cleaned,
                    source_page=page_index,
                    source_meta={
                        "bbox": bbox,
                        "layout_role": self._layout_role(float(y0), float(y1), page_height),
                        "parser_strategy": "pymupdf",
                        "modality": "text",
                        "page_layout": f"page_{page_index}",
                    },
                )
            )
        return nodes

    def _extract_image_regions(self, page, page_index: int, table_bboxes: list[object]) -> list[PdfImageRegion]:
        page_dict = page.get_text("dict")
        regions: list[PdfImageRegion] = []
        order = 0
        page_area = max(float(page.rect.width * page.rect.height), 1.0)
        for block in page_dict.get("blocks", []):
            if block.get("type") != 1:
                continue
            bbox = block.get("bbox")
            if not isinstance(bbox, (list, tuple)) or len(bbox) != 4:
                continue
            normalized_bbox = [float(value) for value in bbox]
            area = max((normalized_bbox[2] - normalized_bbox[0]) * (normalized_bbox[3] - normalized_bbox[1]), 0.0)
            if area / page_area < 0.01:
                continue
            if self._overlaps_any(normalized_bbox, table_bboxes):
                continue
            image_region = PdfImageRegion(
                page_no=page_index,
                bbox=normalized_bbox,
                image_region_id=str(uuid.uuid4()),
                order=order,
            )
            order += 1
            regions.append(image_region)

        if not regions:
            for index, block in enumerate(page.get_text("blocks")):
                x0, y0, x1, y1, text, *_ = block
                normalized_bbox = [float(x0), float(y0), float(x1), float(y1)]
                area = max((normalized_bbox[2] - normalized_bbox[0]) * (normalized_bbox[3] - normalized_bbox[1]), 0.0)
                block_text = (text or "").strip()
                if block_text and not block_text.lower().startswith("<image:"):
                    continue
                if area / page_area < 0.03:
                    continue
                if self._overlaps_any(normalized_bbox, table_bboxes):
                    continue
                regions.append(
                    PdfImageRegion(
                        page_no=page_index,
                        bbox=normalized_bbox,
                        image_region_id=str(uuid.uuid4()),
                        order=index,
                        parser_strategy="pdf_image_region_heuristic",
                    )
                )

        for _ in regions:
            PDF_IMAGE_REGION_DETECTED.inc()
        return regions

    def _extract_image_nodes(
        self,
        page,
        page_index: int,
        filename: str,
        image_regions: list[PdfImageRegion],
    ) -> list[DocumentNode]:
        image_nodes: list[DocumentNode] = []
        for region in image_regions:
            try:
                pixmap = page.get_pixmap(matrix=fitz.Matrix(2, 2), clip=fitz.Rect(*region.bbox), alpha=False)
                region_nodes = self.visual_analyzer.analyze_cropped_region(
                    pixmap.tobytes("png"),
                    f"{filename}-page-{page_index}-region-{region.order}.png",
                    page_no=page_index,
                    bbox=region.bbox,
                    image_region_id=region.image_region_id,
                )
                for node in region_nodes:
                    node.source_meta.setdefault("bbox", region.bbox)
                    node.source_meta.setdefault("layout_role", region.layout_role)
                    node.source_meta.setdefault("image_region_id", region.image_region_id)
                    node.source_meta.setdefault("parser_strategy", "vision_image_region")
                    node.source_meta.setdefault("modality", "image")
                    node.source_meta.setdefault("order", region.order)
                    node.source_meta.setdefault("page_layout", f"page_{page_index}")
                image_nodes.extend(region_nodes)
                PDF_IMAGE_REGION_VISION_SUCCESS.inc()
            except Exception:
                PDF_IMAGE_REGION_VISION_ERROR.inc()
        return image_nodes

    def _assemble_page_nodes(
        self,
        text_nodes: list[DocumentNode],
        table_nodes: list[DocumentNode],
        image_nodes: list[DocumentNode],
    ) -> list[DocumentNode]:
        def sort_key(node: DocumentNode) -> tuple[float, float, int]:
            bbox = node.source_meta.get("bbox")
            if isinstance(bbox, list) and len(bbox) == 4:
                return (float(bbox[1]), float(bbox[0]), int(node.source_meta.get("order") or 0))
            return (10**9, 10**9, int(node.source_meta.get("order") or 0))

        return sorted([*text_nodes, *table_nodes, *image_nodes], key=sort_key)

    def _extract_with_pypdf(self, file_bytes: bytes) -> list[DocumentNode]:
        reader = PdfReader(io.BytesIO(file_bytes))
        nodes: list[DocumentNode] = []
        for page_index, page in enumerate(reader.pages, start=1):
            raw = (page.extract_text() or "").strip()
            if not raw:
                continue
            sections = [section.strip() for section in re.split(r"\n{2,}", raw) if section.strip()]
            for section in sections:
                node_type, level = self._infer_pdf_node_type(section)
                nodes.append(
                    DocumentNode(
                        node_id=str(uuid.uuid4()),
                        node_type=node_type,
                        level=level,
                        text=section,
                        source_page=page_index,
                        source_meta={
                            "parser_strategy": "pypdf",
                            "modality": "text",
                            "page_layout": f"page_{page_index}",
                        },
                    )
                )
        return nodes

    def _extract_tables_from_page(self, page, page_index: int) -> list[DocumentNode]:
        if not hasattr(page, "find_tables"):
            return []
        try:
            tables = page.find_tables()
        except Exception:
            return []
        nodes: list[DocumentNode] = []
        for table in getattr(tables, "tables", []):
            rows = []
            for row in table.extract() or []:
                cells = ["" if cell is None else str(cell).strip() for cell in row]
                if any(cells):
                    rows.append(" | ".join(cells))
            table_text = "\n".join(rows).strip()
            if table_text:
                bbox = getattr(table, "bbox", None)
                nodes.append(
                    DocumentNode(
                        node_id=str(uuid.uuid4()),
                        node_type="table",
                        text=table_text,
                        source_page=page_index,
                        source_meta={
                            "layout_role": "table",
                            "parser_strategy": "pymupdf_table",
                            "bbox": list(bbox) if bbox else None,
                            "modality": "text",
                            "page_layout": f"page_{page_index}",
                        },
                    )
                )
        return nodes

    def _clean_pdf_text(self, text: str) -> str:
        lines = [re.sub(r"\s+", " ", line).strip() for line in text.splitlines()]
        lines = [line for line in lines if line]
        return "\n".join(lines).strip()

    def _infer_pdf_node_type(self, text: str) -> tuple[str, int]:
        stripped = text.strip()
        if "|" in stripped and stripped.count("|") >= 2:
            return "table", 0
        if re.match(r"^\s*(?:[-*+]|\d+[.)])\s+", stripped):
            return "list", 0
        first_line = stripped.splitlines()[0].strip()
        if len(stripped.splitlines()) == 1 and len(first_line) <= 60 and not re.search(r"[。?!;:：；]", first_line):
            return "title", 1
        return "paragraph", 0

    def _layout_role(self, y0: float, y1: float, page_height: float) -> str:
        if y1 <= page_height * 0.12:
            return "header_zone"
        if y0 >= page_height * 0.88:
            return "footer_zone"
        return "body"

    def _sort_blocks_by_columns(self, blocks) -> list[tuple]:
        centers = [((float(item[0]) + float(item[2])) / 2.0) for item in blocks if (item[4] or "").strip()]
        if not centers:
            return []
        median_x = sorted(centers)[len(centers) // 2]

        def key(item):
            x0, y0, x1, _y1, text, *_ = item
            if not (text or "").strip():
                return (99, 99, 99)
            center = (float(x0) + float(x1)) / 2.0
            column = 0 if center <= median_x else 1
            return (column, round(float(y0) / 12), float(x0))

        return sorted(blocks, key=key)

    def _overlaps_any(self, bbox: list[float], other_bboxes: list[object]) -> bool:
        x0, y0, x1, y1 = bbox
        for other in other_bboxes:
            if not isinstance(other, list) or len(other) != 4:
                continue
            ox0, oy0, ox1, oy1 = other
            if x0 < ox1 and x1 > ox0 and y0 < oy1 and y1 > oy0:
                return True
        return False

    def _remove_repeated_page_noise(self, nodes: list[DocumentNode]) -> list[DocumentNode]:
        text_counter: Counter[str] = Counter()
        for node in nodes:
            text = node.text.strip()
            if text and len(text) <= 100:
                text_counter[text] += 1

        repeated_texts = {text for text, count in text_counter.items() if count >= 3}
        cleaned: list[DocumentNode] = []
        for node in nodes:
            text = node.text.strip()
            layout_role = node.source_meta.get("layout_role")
            is_page_number = bool(self.page_number_re.match(text))
            is_repeated_margin = text in repeated_texts and layout_role in {None, "header_zone", "footer_zone"}
            if is_page_number or is_repeated_margin:
                continue
            cleaned.append(node)
        return cleaned


class XlsxParser(BaseParser):
    def parse(self, file_bytes: bytes, filename: str) -> list[DocumentNode]:
        workbook = load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True)
        nodes: list[DocumentNode] = []
        for sheet in workbook.worksheets:
            nodes.append(
                DocumentNode(
                    node_id=str(uuid.uuid4()),
                    node_type="title",
                    level=1,
                    text=sheet.title,
                    source_meta={"sheet_name": sheet.title, "parser_strategy": "xlsx_sheet"},
                )
            )
            rows = []
            for row in sheet.iter_rows(values_only=True):
                values = ["" if value is None else str(value).strip() for value in row]
                if any(values):
                    rows.append(" | ".join(values))
            table_text = "\n".join(rows).strip()
            if table_text:
                nodes.append(
                    DocumentNode(
                        node_id=str(uuid.uuid4()),
                        node_type="table",
                        text=table_text,
                        source_meta={"sheet_name": sheet.title, "parser_strategy": "xlsx_table"},
                    )
                )
        return nodes


class XlsParser(BaseParser):
    def parse(self, file_bytes: bytes, filename: str) -> list[DocumentNode]:
        workbook = xlrd.open_workbook(file_contents=file_bytes)
        nodes: list[DocumentNode] = []
        for sheet in workbook.sheets():
            nodes.append(
                DocumentNode(
                    node_id=str(uuid.uuid4()),
                    node_type="title",
                    level=1,
                    text=sheet.name,
                    source_meta={"sheet_name": sheet.name, "parser_strategy": "xls_sheet"},
                )
            )
            rows = []
            for row_index in range(sheet.nrows):
                values = []
                for value in sheet.row_values(row_index):
                    text = "" if value is None else str(value).strip()
                    if text.endswith(".0"):
                        text = text[:-2]
                    values.append(text)
                if any(values):
                    rows.append(" | ".join(values))
            table_text = "\n".join(rows).strip()
            if table_text:
                nodes.append(
                    DocumentNode(
                        node_id=str(uuid.uuid4()),
                        node_type="table",
                        text=table_text,
                        source_meta={"sheet_name": sheet.name, "parser_strategy": "xls_table"},
                    )
                )
        return nodes


def get_parser(filename: str) -> BaseParser:
    lower = filename.lower()
    if lower.endswith(".docx"):
        return DocxParser()
    if lower.endswith(".pdf"):
        return PdfParser()
    if lower.endswith(".xlsx"):
        return XlsxParser()
    if lower.endswith(".xls"):
        return XlsParser()
    if lower.endswith(".txt") or lower.endswith(".md"):
        return TxtMarkdownParser()
    raise ValueError(f"Unsupported file type for {filename}")


def is_image_filename(filename: str) -> bool:
    return Path(filename).suffix.lower() in IMAGE_SUFFIXES
